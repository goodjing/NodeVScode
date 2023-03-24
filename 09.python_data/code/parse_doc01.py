import docx

# 1. python-docx提取文字
# doc = docx.Document(r'E:\NodeVScode\09.python_data\EXData\data05doc.docx')
# print(doc.paragraphs)
# for paragraph in doc.paragraphs:
#     print(paragraph.text)

# 2. python-docx提取文字块
# doc = docx.Document(r'E:\NodeVScode\09.python_data\EXData\data05doc.docx')
# print(doc.paragraphs)
# paragraph = doc.paragraphs[1]  # 定位到第二段
# runs = paragraph.runs
# print(runs)
# for run in paragraph.runs:
#     print(run.text)

# 3. 在Word文档中添加段落
# doc = docx.Document(r'E:\NodeVScode\09.python_data\EXData\data05doc.docx')
# paragraph1 = doc.add_paragraph('这是新增的段落')
# doc.save(r'E:\NodeVScode\09.python_data\EXData\data05doc.docx')

# 4. 添加文字块
# doc = docx.Document(r'E:\NodeVScode\09.python_data\EXData\data05doc.docx')
# paragraph3 = doc.add_paragraph()
# paragraph3.add_run('我被加粗了文字块儿').bold = True
# paragraph3.add_run('我是普通文字块儿')
# paragraph3.add_run('我是斜体字块').italic = True
# doc.save(r'E:\NodeVScode\09.python_data\EXData\data05doc.docx')


# 5. 添加一个分页
doc = docx.Document(r'E:\NodeVScode\09.python_data\EXData\data05doc.docx')
doc.add_page_break()  # 添加一个分页
doc.save(r'E:\NodeVScode\09.python_data\EXData\data05doc.docx')